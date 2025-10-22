import os
import streamlit as st
import tempfile
from moviepy.editor import ImageClip, concatenate_videoclips, AudioFileClip
from fpdf import FPDF
from docx import Document
from PIL import Image, ImageDraw
import pyttsx3
from PyPDF2 import PdfReader
from moviepy.video.io.VideoFileClip import VideoFileClip

# ---------- CONFIG ----------
st.set_page_config(page_title="🧠 Offline AI Assistant", layout="wide")
st.title("🧠 Offline AI Assistant (Offline)")

BASE_DIR = os.path.dirname(__file__)
GENERATED_DIR = os.path.join(BASE_DIR, "generated")
os.makedirs(GENERATED_DIR, exist_ok=True)

for folder in ["pdf", "docx", "txt", "image", "video", "audio", "captions"]:
    os.makedirs(os.path.join(GENERATED_DIR, folder), exist_ok=True)

# ---------- Helper Functions ----------
def read_file_content(uploaded_file):
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "txt":
        return uploaded_file.read().decode("utf-8")
    elif ext == "pdf":
        reader = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif ext == "docx":
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)
    return ""

def text_to_speech(text, audio_path):
    engine = pyttsx3.init()
    engine.save_to_file(text, audio_path)
    engine.runAndWait()

def text_to_video_line_clips(text, output_path, clip_size=(1280, 720), line_duration=3):
    """Generate video from text, each line as a separate clip with TTS audio"""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        lines = [text[:200]]

    clips = []
    audio_clips = []

    for i, line in enumerate(lines):
        # Generate audio
        audio_path = os.path.join(tempfile.gettempdir(), f"temp_audio_{i}.mp3")
        text_to_speech(line, audio_path)
        audio_clip = AudioFileClip(audio_path)
        audio_clips.append(audio_clip)

        # Generate image clip
        img = Image.new("RGB", clip_size, color="black")
        draw = ImageDraw.Draw(img)
        draw.text((50, clip_size[1]//2 - 25), line[:200], fill="white")
        temp_img_path = os.path.join(tempfile.gettempdir(), f"temp_img_{i}.png")
        img.save(temp_img_path)
        img_clip = ImageClip(temp_img_path).set_duration(line_duration).set_audio(audio_clip)
        clips.append(img_clip)

    final_clip = concatenate_videoclips(clips, method="compose")
    final_clip.write_videofile(output_path, fps=24, codec="libx264", audio_codec="aac", logger=None)

    # Cleanup temp files
    for i in range(len(lines)):
        img_file = os.path.join(tempfile.gettempdir(), f"temp_img_{i}.png")
        audio_file = os.path.join(tempfile.gettempdir(), f"temp_audio_{i}.mp3")
        if os.path.exists(img_file):
            os.remove(img_file)
        if os.path.exists(audio_file):
            os.remove(audio_file)

# ---------- SIDEBAR ----------
tabs = st.sidebar.radio("📍 Choose Section", [
    "📘 Summarizer",
    "🧠 Semantic Analyzer",
    "🎨 File Generator",
    "🎬 Video Caption Generator"
])

# ---------- SUMMARIZER ----------
if tabs == "📘 Summarizer":
    st.header("📘 Text & File Summarizer (Offline)")
    uploaded_file = st.file_uploader("Upload a file (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], key="sum_upload")
    text_input = st.text_area("Or paste text here:")

    content = read_file_content(uploaded_file) if uploaded_file else text_input

    if st.button("Summarize"):
        if content:
            sentences = content.split(". ")
            summary = ". ".join(sentences[:5]) + "..." if len(sentences) > 5 else content
            st.subheader("🧾 Summary:")
            st.write(summary)
        else:
            st.warning("Provide text or upload a file.")

# ---------- SEMANTIC ANALYZER ----------
elif tabs == "🧠 Semantic Analyzer":
    st.header("🧠 Semantic Analyzer (Offline)")
    uploaded_file = st.file_uploader("Upload a file (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], key="sem_upload")
    text_input = st.text_area("Or paste text here:")

    content = read_file_content(uploaded_file) if uploaded_file else text_input

    if st.button("Analyze"):
        if content:
            keywords = [word for word in content.split() if len(word) > 5]
            st.subheader("🔍 Keywords:")
            st.write(", ".join(keywords[:10]))
            st.write(f"📏 Word count: {len(content.split())}")
        else:
            st.warning("Provide text or upload a file.")

# ---------- FILE GENERATOR ----------
elif tabs == "🎨 File Generator":
    st.header("🎨 File Generator (Choose type to generate)")
    content = st.text_area("Enter text content:")

    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        if st.button("Generate TXT"):
            if content.strip():
                txt_path = os.path.join(GENERATED_DIR, "txt", "generated.txt")
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(content)
                st.success("✅ TXT generated")
                st.download_button("⬇️ Download TXT", open(txt_path, "rb"), "generated.txt")
            else:
                st.warning("Enter some text first")
    with col2:
        if st.button("Generate DOCX"):
            if content.strip():
                docx_path = os.path.join(GENERATED_DIR, "docx", "generated.docx")
                doc = Document()
                doc.add_paragraph(content)
                doc.save(docx_path)
                st.success("✅ DOCX generated")
                st.download_button("⬇️ Download DOCX", open(docx_path, "rb"), "generated.docx")
            else:
                st.warning("Enter some text first")
    with col3:
        if st.button("Generate PDF"):
            if content.strip():
                pdf_path = os.path.join(GENERATED_DIR, "pdf", "generated.pdf")
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, content)
                pdf.output(pdf_path)
                st.success("✅ PDF generated")
                st.download_button("⬇️ Download PDF", open(pdf_path, "rb"), "generated.pdf")
            else:
                st.warning("Enter some text first")
    with col4:
        if st.button("Generate Image"):
            if content.strip():
                img_path = os.path.join(GENERATED_DIR, "image", "generated.png")
                img = Image.new("RGB", (800, 400), color="white")
                draw = ImageDraw.Draw(img)
                draw.text((10, 10), content[:400], fill="black")
                img.save(img_path)
                st.success("✅ Image generated")
                st.image(img_path)
                st.download_button("⬇️ Download Image", open(img_path, "rb"), "generated.png")
            else:
                st.warning("Enter some text first")
    with col5:
        if st.button("Generate Video"):
            if content.strip():
                output_video_path = os.path.join(GENERATED_DIR, "video", "generated_text_video.mp4")
                st.info("Generating video from text with audio...")
                from moviepy.editor import AudioFileClip  # ensure fresh import
                def generate_video():
                    lines = [line.strip() for line in content.split("\n") if line.strip()]
                    if not lines:
                        lines = [content[:200]]
                    clips = []
                    for i, line in enumerate(lines):
                        audio_path = os.path.join(tempfile.gettempdir(), f"temp_audio_{i}.mp3")
                        text_to_speech(line, audio_path)
                        audio_clip = AudioFileClip(audio_path)
                        img = Image.new("RGB", (1280, 720), color="black")
                        draw = ImageDraw.Draw(img)
                        draw.text((50, 350), line[:200], fill="white")
                        temp_img_path = os.path.join(tempfile.gettempdir(), f"temp_img_{i}.png")
                        img.save(temp_img_path)
                        img_clip = ImageClip(temp_img_path).set_duration(3).set_audio(audio_clip)
                        clips.append(img_clip)
                    final_clip = concatenate_videoclips(clips, method="compose")
                    final_clip.write_videofile(output_video_path, fps=24, codec="libx264", audio_codec="aac", logger=None)
                generate_video()
                st.success("✅ Video generated")
                st.video(output_video_path)
                st.download_button("⬇️ Download Video with Audio", open(output_video_path, "rb"), "generated_text_video.mp4")
            else:
                st.warning("Enter some text first")

# ---------- VIDEO CAPTION GENERATOR ----------
elif tabs == "🎬 Video Caption Generator":
    import whisper
    import subprocess

    st.set_page_config(page_title="Offline AI Video Caption Generator", layout="wide")
    st.title("🎬 Offline AI Video Caption Generator (Video + Captions)")

    # -------------------------------
    # Project-scoped FFmpeg setup
    # -------------------------------
    ffmpeg_bin = os.path.join(os.path.dirname(__file__), "ffmpeg", "bin", "ffmpeg.exe")
    os.environ["IMAGEIO_FFMPEG_EXE"] = ffmpeg_bin
    os.environ["PATH"] += os.pathsep + os.path.dirname(ffmpeg_bin)

    # -------------------------------
    # File uploader
    # -------------------------------
    video_file = st.file_uploader(
        "Upload Video (MP4, MOV, AVI, MKV)", type=["mp4", "mov", "avi", "mkv"]
    )

    # -------------------------------
    # Choose caption language
    # -------------------------------
    caption_option = st.radio(
        "Choose Caption Language:",
        ("Original Language", "English Translation")
    )


    # -------------------------------
    # Helper function for SRT
    # -------------------------------
    def format_timestamp(seconds):
        milliseconds = int((seconds - int(seconds)) * 1000)
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02}:{minutes:02}:{secs:02},{milliseconds:03}"


    if video_file is not None:
        st.video(video_file)

        # Save uploaded video in project folder
        video_path = os.path.join(os.path.dirname(__file__), video_file.name)
        with open(video_path, "wb") as f:
            f.write(video_file.read())

        if st.button("Generate Captions and Burn into Video"):
            st.info("Loading Whisper model...")
            model = whisper.load_model("small")

            st.write("🎧 Extracting audio from video...")
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio:
                audio_path = temp_audio.name
                clip = VideoFileClip(video_path)
                clip.audio.write_audiofile(audio_path, codec='pcm_s16le', logger=None)
                clip.close()

            st.write("📝 Transcribing audio...")
            task_type = "translate" if caption_option == "English Translation" else "transcribe"
            result = model.transcribe(audio_path, verbose=False, task=task_type)

            # -------------------------------
            # Save captions as SRT
            # -------------------------------
            suffix = "_english" if task_type == "translate" else "_original"
            srt_path = os.path.splitext(video_path)[0] + f"{suffix}.srt"
            with open(srt_path, "w", encoding="utf-8") as f:
                for i, seg in enumerate(result["segments"], start=1):
                    start = format_timestamp(seg["start"])
                    end = format_timestamp(seg["end"])
                    text = seg["text"].strip()
                    f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

            st.success("✅ Captions generated!")

            # -------------------------------
            # Burn captions into video
            # -------------------------------
            st.write("🎞️ Burning captions into video...")
            output_video_path = os.path.splitext(video_path)[0] + f"{suffix}_subtitled.mp4"

            burn_cmd = [
                ffmpeg_bin,
                "-y",
                "-i", video_path,
                "-vf", f"subtitles={srt_path}",
                output_video_path
            ]
            subprocess.run(burn_cmd, check=True)

            st.success("✅ Video with captions created!")

            # -------------------------------
            # Download buttons
            # -------------------------------
            st.download_button(
                "⬇️ Download Captions (.srt)",
                open(srt_path, "rb"),
                file_name=os.path.basename(srt_path)
            )

            st.download_button(
                "⬇️ Download Video with Subtitles (.mp4)",
                open(output_video_path, "rb"),
                file_name=os.path.basename(output_video_path)
            )

            # -------------------------------
            # Previews
            # -------------------------------
            st.write("📋 **Preview of captions (first few lines):**")
            for seg in result["segments"][:5]:
                st.write(f"**{format_timestamp(seg['start'])} → {format_timestamp(seg['end'])}** — {seg['text']}")

            st.write("🎞️ **Preview Video with Captions:**")
            st.video(output_video_path)
